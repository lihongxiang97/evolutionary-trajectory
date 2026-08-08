from pathlib import Path
import hashlib
import re
import zipfile

from docx import Document
from openpyxl import load_workbook


ROOT = Path("work/plant_communications_submission_20260805")
MAIN = ROOT / "Manuscript_Plant_Communications.docx"
SOURCE = Path("work/submission_ready_figure_fixed_20260804/Manuscript_T2T_hapA_hapB_3D_mechanisms.docx")
SUPP = ROOT / "Supplementary_Information.docx"
BOOK = ROOT / "Supplementary_Tables.xlsx"


checks = []


def check(label, condition, detail=""):
    checks.append((label, bool(condition), detail))


for label, path in [("main DOCX ZIP integrity", MAIN), ("supplement DOCX ZIP integrity", SUPP)]:
    try:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
        check(label, bad is None, str(bad or "OK"))
    except Exception as exc:
        check(label, False, str(exc))

main = Document(MAIN)
source = Document(SOURCE)
supp = Document(SUPP)
text = "\n".join(x.text for x in main.paragraphs)

check("main manuscript contains 8 embedded figures", len(main.inline_shapes) == 8, str(len(main.inline_shapes)))
check("supplement contains 18 embedded figures", len(supp.inline_shapes) == 18, str(len(supp.inline_shapes)))
check("Results and figure legends are unchanged", all(source.paragraphs[i].text == main.paragraphs[i].text for i in range(19, 83)))
check("single-paragraph Summary is 150-250 words", 150 <= len(main.paragraphs[4].text.split()) <= 250, str(len(main.paragraphs[4].text.split())))
check("BMC structured-summary labels removed", not any(x in text for x in ("Background:", "Results:", "Conclusions:")))
check("journal-facing section headings present", all(x in text for x in ("Summary", "Introduction", "Materials and methods", "Resource availability")))

main_figs = set(map(int, re.findall(r"(?<!Supplementary )Fig[s]?\. (?:S[0-9]+(?: and S[0-9]+)*|[0-9]+)", text))) if False else set(map(int, re.findall(r"(?<!Supplementary )Fig\. ([0-9]+)", text)))
supp_figs = set(map(int, re.findall(r"Supplementary Fig(?:s)?\. S([0-9]+)", text)))
for group in re.findall(r"Supplementary Figs\. ([^.;\n]+)", text):
    supp_figs.update(map(int, re.findall(r"S([0-9]+)", group)))
tables = set(map(int, re.findall(r"Supplementary Table ([0-9]+)", text)))
check("all main figures cited", main_figs == set(range(1, 9)), str(sorted(main_figs)))
check("all supplementary figures cited", supp_figs == set(range(1, 19)), str(sorted(supp_figs)))
check("all supplementary tables cited", tables == set(range(1, 22)), str(sorted(tables)))

forbidden = ["reanalysis", "old genome", "unable to prove", "cannot prove"]
hits = [x for x in forbidden if x in text.lower()]
check("no workflow-history or nihilistic wording", not hits, ", ".join(hits) or "none")

wb = load_workbook(BOOK, read_only=True, data_only=True)
expected_sheets = {"Contents"} | {f"Table {i}" for i in range(1, 22)}
check("workbook contains Contents and Table 1-21", set(wb.sheetnames) == expected_sheets, str(wb.sheetnames))
check("no letter-suffixed supplementary tables", not any(re.fullmatch(r"Table [0-9]+[a-z]", s, re.I) for s in wb.sheetnames))

highlights = [x.strip() for x in (ROOT / "Highlights.txt").read_text(encoding="utf-8").splitlines() if x.strip()]
check("four Highlights supplied", len(highlights) == 4, str(len(highlights)))
check("every Highlight is <=85 characters", all(len(x) <= 85 for x in highlights), str([len(x) for x in highlights]))

formats = {".pdf", ".svg", ".png", ".tiff"}
for folder, expected in [(ROOT / "Figures_Main", 8), (ROOT / "Figures_Supplementary", 18)]:
    stems = {}
    for f in folder.iterdir():
        if f.is_file() and f.suffix.lower() in formats:
            stems.setdefault(f.stem, set()).add(f.suffix.lower())
    check(f"{folder.name} has {expected} figures in four formats", len(stems) == expected and all(v == formats for v in stems.values()), f"figures={len(stems)}")

admin = []
if not main.paragraphs[181].text.strip():
    admin.append("Funding statement")
if not main.paragraphs[183].text.strip():
    admin.append("Authors' contributions")
admin.extend(["author names and affiliations", "corresponding-author email", "stable public code release/DOI"])

overall = all(ok for _, ok, _ in checks)
lines = ["# Plant Communications submission content audit", "", f"Overall content status: **{'PASS' if overall else 'FAIL'}**", ""]
for label, ok, detail in checks:
    lines.append(f"- {'PASS' if ok else 'FAIL'} - {label}" + (f" ({detail})" if detail else ""))
lines += ["", "## Administrative items requiring author input", ""] + [f"- {x}" for x in admin]
(ROOT / "Reports" / "CONTENT_AUDIT_PLANT_COMMUNICATIONS.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

# Package hashes exclude the checksum file itself.
rows = ["SHA256\tFile"]
for f in sorted((x for x in ROOT.rglob("*") if x.is_file() and x.name != "SHA256SUMS.tsv"), key=lambda x: x.as_posix().lower()):
    digest = hashlib.sha256(f.read_bytes()).hexdigest()
    rows.append(f"{digest}\t{f.relative_to(ROOT).as_posix()}")
(ROOT / "SHA256SUMS.tsv").write_text("\n".join(rows) + "\n", encoding="utf-8")

print("PASS" if overall else "FAIL")
for label, ok, detail in checks:
    print(f"{'PASS' if ok else 'FAIL'}\t{label}\t{detail}")
